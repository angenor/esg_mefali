"""Export des dossiers de candidature en PDF (WeasyPrint) et Word (python-docx)."""

import io
import logging
import re
from datetime import datetime, timezone
from pathlib import Path

from jinja2 import Environment, FileSystemLoader

logger = logging.getLogger(__name__)

TEMPLATE_DIR = Path(__file__).parent / "templates"

TARGET_TYPE_LABELS = {
    "fund_direct": "Candidature directe",
    "intermediary_bank": "Via banque partenaire",
    "intermediary_agency": "Via agence d'implémentation",
    "intermediary_developer": "Via développeur carbone",
}

SECTION_STATUS_LABELS = {
    "not_generated": "Non rédigée",
    "generated": "Générée",
    "validated": "Validée",
}


def _prepare_sections(sections: dict) -> list[dict]:
    """Preparer les sections pour le template."""
    result = []
    for key, section in sections.items():
        result.append({
            "key": key,
            "title": section.get("title", key),
            "content": section.get("content"),
            "status": section.get("status", "not_generated"),
            "status_label": SECTION_STATUS_LABELS.get(
                section.get("status", "not_generated"), "Inconnu"
            ),
        })
    return result


def _render_export_html(application) -> str:
    """Rendre le template HTML pour l'export."""
    env = Environment(
        loader=FileSystemLoader(str(TEMPLATE_DIR)),
        autoescape=True,
    )
    template = env.get_template("application_export.html")

    target_val = application.target_type.value if hasattr(application.target_type, 'value') else application.target_type

    return template.render(
        fund_name=application.fund.name if application.fund else "Fonds",
        fund_organization=application.fund.organization if application.fund else "",
        intermediary_name=application.intermediary.name if application.intermediary else None,
        target_type_label=TARGET_TYPE_LABELS.get(target_val, target_val),
        generated_date=datetime.now(timezone.utc).strftime("%d/%m/%Y"),
        sections=_prepare_sections(application.sections or {}),
    )


def _strip_html(html: str) -> str:
    """Retirer les balises HTML pour le contenu Word."""
    text = re.sub(r"<br\s*/?>", "\n", html)
    text = re.sub(r"<li>", "• ", text)
    text = re.sub(r"</li>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _export_docx(application) -> bytes:
    """Generer le document Word."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    doc = Document()

    # Styles
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Titre
    title = doc.add_heading("Dossier de Candidature", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    fund_name = application.fund.name if application.fund else "Fonds"
    fund_org = application.fund.organization if application.fund else ""
    subtitle = doc.add_paragraph(f"{fund_name} — {fund_org}")
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if application.intermediary:
        inter_para = doc.add_paragraph(f"Via {application.intermediary.name}")
        inter_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("")  # Espacement

    # Sections
    for key, section in (application.sections or {}).items():
        section_title = section.get("title", key)
        doc.add_heading(section_title, level=1)

        content = section.get("content")
        if content:
            plain_text = _strip_html(content)
            for paragraph_text in plain_text.split("\n\n"):
                if paragraph_text.strip():
                    doc.add_paragraph(paragraph_text.strip())
        else:
            para = doc.add_paragraph("Section non encore rédigée.")
            para.runs[0].font.italic = True
            para.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph(
        "Dossier généré par ESG Mefali — Plateforme de finance durable pour les PME africaines"
    )
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


async def export_application(
    application,
    format: str = "pdf",
) -> tuple[bytes, str, str]:
    """Exporter le dossier en PDF ou Word.

    Retourne (bytes, content_type, filename).
    """
    fund_name = (application.fund.name if application.fund else "dossier").lower().replace(" ", "_")

    if format == "pdf":
        html_content = _render_export_html(application)
        from weasyprint import HTML
        pdf_bytes = HTML(string=html_content).write_pdf()
        return pdf_bytes, "application/pdf", f"dossier_{fund_name}.pdf"

    if format == "docx":
        docx_bytes = _export_docx(application)
        return (
            docx_bytes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"dossier_{fund_name}.docx",
        )

    raise ValueError(f"Format non supporte : {format}. Utiliser 'pdf' ou 'docx'.")
